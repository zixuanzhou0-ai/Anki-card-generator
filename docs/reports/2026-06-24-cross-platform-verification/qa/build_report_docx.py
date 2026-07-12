from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_ROOT = Path(r"E:\ANKI\docs\reports\2026-06-24-cross-platform-verification")
SCREENSHOT_DIR = Path(r"E:\ANKI\docs\screenshots")
OUT = REPORT_ROOT / "Anki_Card_Generator_Cross_Platform_Verification.docx"

FONT = "Microsoft YaHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(92, 105, 121)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
PASS = RGBColor(21, 107, 74)
WARN = RGBColor(122, 90, 0)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size=10.5, color=INK, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph_border_bottom(paragraph, color="2E74B5", size="12", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_para(doc, text="", style=None, size=10.5, color=INK, bold=None, italic=None, align=None, before=0, after=6, line=1.1):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        set_run_font(r, size=16, color=BLUE, bold=True)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_run_font(r, size=13, color=BLUE, bold=True)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
    return p


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1800, 2520, 2520, 2520]
    set_table_geometry(table, widths)
    headers = ["Area", "Evidence", "Result", "Notes"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(value)
            color = PASS if value == "PASS" else WARN if value == "LIMITED" else INK
            set_run_font(r, size=9.2, color=color, bold=(i == 2))
    add_para(doc, "", after=4)
    return table


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT_GRAY)
        cells[0].text = ""
        cells[1].text = ""
        r1 = cells[0].paragraphs[0].add_run(label)
        r2 = cells[1].paragraphs[0].add_run(value)
        set_run_font(r1, size=9.3, color=INK, bold=True)
        set_run_font(r2, size=9.3, color=INK)
    add_para(doc, "", after=4)
    return table


def add_callout(doc, title, text, status="PASS"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    title_run = p.add_run(f"{title}: ")
    set_run_font(title_run, size=10.5, color=PASS if status == "PASS" else WARN, bold=True)
    body_run = p.add_run(text)
    set_run_font(body_run, size=10.2, color=INK)
    add_para(doc, "", after=2)


def add_picture_if_exists(doc, path, caption, width=6.1):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        cap = add_para(doc, caption, size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
        return cap
    add_callout(doc, "Missing screenshot", str(path), status="LIMITED")
    return None


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10.5)

    for style_name in ["Header", "Footer"]:
        styles[style_name].font.name = FONT
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        styles[style_name].font.size = Pt(8.5)
        styles[style_name].font.color.rgb = MUTED

    header = section.header.paragraphs[0]
    header.text = ""
    r = header.add_run("Anki Card Generator | v0.9.4-beta Verification")
    set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Confidential local verification artifact")
    set_run_font(r, size=8.5, color=MUTED)
    return doc


def build():
    doc = setup_document()

    add_para(doc, "VERIFICATION MEMO", size=11, color=BLUE, bold=True, after=2)
    title = add_para(
        doc,
        "Anki Card Generator",
        size=24,
        color=INK,
        bold=True,
        after=3,
    )
    subtitle = add_para(
        doc,
        "v0.9.4-beta GitHub Release, Windows Desktop Validation, and Browser Helper Roadmap",
        size=13,
        color=MUTED,
        after=14,
    )
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="2E74B5", size="14", space="4")

    add_two_col_table(
        doc,
        [
            ("Prepared for", "Anki Card Generator product owner and open-source users"),
            ("Date", "2026-06-24"),
            ("Release under review", "v0.9.4-beta"),
            ("Primary platform", "Windows desktop"),
            ("Browser status", "Planned as a separate Browser + Local Helper product line"),
            ("Evidence policy", "No API keys, APKG/media artifacts, or private local data are intended for GitHub"),
        ],
    )

    add_callout(
        doc,
        "Executive conclusion",
        "The Windows desktop release is usable for public beta testing. GitHub labeling, release assets, hash verification, automated UI checks, worker tests, Tauri packaging, and APKG smoke verification passed. The Browser + Local Helper work should remain a separate product line and must not be mixed into the current desktop folder.",
    )

    add_heading(doc, "1. Verification Snapshot")
    add_status_table(
        doc,
        [
            ("GitHub About", "Repo description labels Windows desktop video-to-Anki generator", "PASS", "Users can understand current platform scope."),
            ("Release assets", "NSIS, MSI, portable zip, SHA256SUMS", "PASS", "All assets are marked as Windows downloads."),
            ("Hash integrity", "Setup exe hash matches GitHub digest and SHA256SUMS", "PASS", "SHA256 61b44fe7...a1637db."),
            ("Installed app", "winget/ARP shows Anki Card Generator 0.9.4", "PASS", "Installed app starts from LocalAppData."),
            ("UI reachability", "Playwright compact-mode smoke", "PASS", "Minimum-size left panel controls are reachable."),
            ("Installed UI CDP", "Production WebView2 did not expose debug port", "LIMITED", "Recorded as observability limit, not product failure."),
            ("Release smoke", "Synthetic video + SRT to APKG verify", "PASS", "Media, audio, TTS, MP4/WebM and fields verified."),
        ],
    )

    add_heading(doc, "2. Public GitHub Readiness")
    add_para(
        doc,
        "The public repository currently tells users that this is a Windows desktop video-to-Anki generator. The release page exposes Windows-specific assets and a checksum file. README screenshots show the workspace, workflow, settings, TTS, environment checks, and final Anki card results.",
    )
    add_picture_if_exists(doc, SCREENSHOT_DIR / "desktop-workspace.png", "Figure 1. Current desktop workspace shown in public documentation.")

    add_heading(doc, "3. Functional Coverage")
    add_status_table(
        doc,
        [
            ("Local video + SRT", "Worker tests and release smoke", "PASS", "Synthetic 8-second video produced APKG."),
            ("Video URL mode", "Playwright workflow shell", "PASS", "UI generation path exercised."),
            ("Learning-point review", "Playwright workflow shell", "PASS", "Review panel appears before APKG generation."),
            ("TTS export", "Worker tests and release smoke", "PASS", "Cached TTS audio is included and referenced."),
            ("Video media", "APKG verify", "PASS", "MP4, WebM, poster and audio fields exist."),
            ("Anki verify", "Worker tests and APKG verify script", "PASS", "No missing or invalid archive media."),
            ("Compact UI", "Playwright 1180x780 smoke", "PASS", "Batch folder picker and bottom CTA stay reachable."),
        ],
    )
    add_picture_if_exists(doc, SCREENSHOT_DIR / "workflow-generated.png", "Figure 2. Review/export stage after cards are generated.")

    add_heading(doc, "4. Release Smoke Evidence")
    add_two_col_table(
        doc,
        [
            ("Command", "npm.cmd run smoke:release"),
            ("Result", "Smoke test passed"),
            ("Segments", "1"),
            ("APKG verify mode", "sqlite_fallback"),
            ("Notes/cards", "1 note / 1 card"),
            ("Media checks", "MP4, WebM, poster, audio, TTS and phrase audio present"),
            ("Missing media", "None"),
            ("Unreferenced media", "None"),
        ],
    )
    add_para(
        doc,
        "The smoke run uses generated test media and a fake cached TTS setup. It proves the product pipeline, packaging, and APKG verification logic without exposing a real model API key.",
    )

    add_heading(doc, "5. Screenshots for Users")
    add_para(
        doc,
        "The public screenshot set already covers the primary user journey. The next screenshot to add should be the installed compact-mode source panel, but the installed WebView2 process did not expose CDP during this run. Existing Playwright screenshots still cover the same reachability behavior in the application shell.",
    )
    add_picture_if_exists(doc, SCREENSHOT_DIR / "settings-model-api.png", "Figure 3. Model and API configuration surface.")
    add_picture_if_exists(doc, SCREENSHOT_DIR / "anki-card-stress-middle.jpg", "Figure 4. Final Anki card example included in public documentation.", width=5.2)

    add_heading(doc, "6. Windows Desktop vs. Browser + Helper")
    add_para(
        doc,
        "The Windows desktop app and the future browser app must remain separate. Tauri gives the existing desktop frontend native backend support; a static browser app does not automatically inherit local file, ffmpeg, worker, APKG export, or AnkiConnect privileges. A local helper is therefore the correct bridge for browser users.",
    )
    add_status_table(
        doc,
        [
            ("Windows desktop", "Tauri app + worker", "PASS", "Current shipped product for Windows users."),
            ("Browser web", "Static UI and local settings", "LIMITED", "Planned; should live in its own app folder."),
            ("Local helper", "Localhost API for native actions", "LIMITED", "Planned; owns files, ffmpeg, APKG and verify."),
            ("Shared packages", "Types, card schema and UI tokens", "LIMITED", "Keep shared code narrow to prevent code sprawl."),
        ],
    )

    add_heading(doc, "7. Tooling and Plugin Use")
    add_two_col_table(
        doc,
        [
            ("GitHub / gh CLI", "Release, asset, README, tag, digest and repository checks."),
            ("Playwright", "UI smoke and compact-mode reachability validation."),
            ("documents:documents", "This Word report and render-based visual QA."),
            ("presentations:Presentations", "PowerPoint deck via @oai/artifact-tool."),
            ("PowerShell / winget", "Install registration, process, hash and local release checks."),
            ("Secret scan", "API key, token, local artifact and diff safety checks."),
        ],
    )

    add_heading(doc, "8. Risks and Required Follow-up")
    add_status_table(
        doc,
        [
            ("CDP on installed app", "WebView2 did not expose debug port", "LIMITED", "Use non-CDP screenshot method next time."),
            ("Full matrix", "27-hour matrix not rerun", "LIMITED", "Reasonable because this report validates release state, not every material permutation."),
            ("Browser product", "Not implemented yet", "LIMITED", "Documented as roadmap only."),
            ("Git status", "Cargo.toml newline warning", "LIMITED", "No content diff; do not stage accidentally."),
            ("Bundle size", "Vite chunk warning", "LIMITED", "Performance optimization candidate, not a release blocker."),
        ],
    )

    add_callout(
        doc,
        "Final recommendation",
        "Keep v0.9.4-beta available as the Windows beta. Publish clearer README routing for Windows Desktop versus Browser + Local Helper. Do not upload local APKG/media evidence. Continue browser work in a separate folder with a narrow shared package layer.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    REPORT_ROOT.mkdir(parents=True, exist_ok=True)
    build()


